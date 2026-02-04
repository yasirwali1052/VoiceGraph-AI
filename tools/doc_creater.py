from docx import Document
from docx.shared import Pt
import os
from config import OUTPUT_DIR

def create_document(content, filename="output.docx"):
    doc = Document()
    
    title = doc.add_heading('Research Report', 0)
    
    for line in content.split('\n'):
        if line.strip():
            doc.add_paragraph(line.strip())
    
    filepath = os.path.join(OUTPUT_DIR, filename)
    doc.save(filepath)
    
    return filepath